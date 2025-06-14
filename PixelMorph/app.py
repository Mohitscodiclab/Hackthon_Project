from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from docx.shared import Inches
import svgwrite
import tempfile
import shutil
import os
import io
import base64
from PIL import Image
import uuid
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Create directories if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class ImageConverter:
    def __init__(self, image_path):
        self.image = Image.open(image_path)
        self.image_path = image_path
    
    def convert_to_image(self, save_path, format_type):
        if format_type == "JPEG" and self.image.mode in ("RGBA", "P"):
            background = Image.new("RGB", self.image.size, (255, 255, 255))
            background.paste(self.image, mask=self.image.split()[-1] if self.image.mode == "RGBA" else None)
            background.save(save_path, "JPEG", quality=95)
        else:
            self.image.save(save_path, format_type)
    
    def convert_to_svg(self, save_path):
        img_buffer = io.BytesIO()
        self.image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
        
        dwg = svgwrite.Drawing(save_path, size=self.image.size)
        dwg.add(dwg.image(
            href=f"data:image/png;base64,{img_base64}",
            insert=(0, 0),
            size=self.image.size
        ))
        dwg.save()
    
    def convert_to_pdf(self, save_path):
        img_buffer = io.BytesIO()
        self.image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        c = canvas.Canvas(save_path, pagesize=letter)
        page_width, page_height = letter
        
        img_width, img_height = self.image.size
        scale = min(page_width / img_width, page_height / img_height) * 0.8
        
        scaled_width = img_width * scale
        scaled_height = img_height * scale
        
        x = (page_width - scaled_width) / 2
        y = (page_height - scaled_height) / 2
        
        c.drawInlineImage(img_buffer, x, y, width=scaled_width, height=scaled_height)
        c.save()
    
    def convert_to_docx(self, save_path):
        img_buffer = io.BytesIO()
        self.image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        doc = Document()
        doc.add_heading('Converted Image', 0)
        
        img_width, img_height = self.image.size
        max_width = 6.0
        
        if img_width > img_height:
            width = Inches(max_width)
            height = Inches((img_height / img_width) * max_width)
        else:
            height = Inches(max_width)
            width = Inches((img_width / img_height) * max_width)
        
        paragraph = doc.add_paragraph()
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(img_buffer, width=width, height=height)
        paragraph.alignment = 1
        
        doc.save(save_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        # Generate unique filename
        filename = str(uuid.uuid4()) + '.' + file.filename.rsplit('.', 1)[1].lower()
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Get image info
        try:
            with Image.open(filepath) as img:
                image_info = {
                    'filename': filename,
                    'original_name': file.filename,
                    'size': img.size,
                    'format': img.format,
                    'mode': img.mode
                }
            return jsonify(image_info)
        except Exception as e:
            os.remove(filepath)
            return jsonify({'error': f'Invalid image file: {str(e)}'}), 400
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/convert', methods=['POST'])
def convert_image():
    data = request.get_json()
    filename = data.get('filename')
    format_type = data.get('format')
    
    if not filename or not format_type:
        return jsonify({'error': 'Missing filename or format'}), 400
    
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(input_path):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        converter = ImageConverter(input_path)
        
        # Generate output filename
        output_filename = f"{filename.rsplit('.', 1)[0]}.{format_type.lower()}"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        # Convert based on format
        if format_type.upper() in ['JPG', 'JPEG']:
            converter.convert_to_image(output_path, 'JPEG')
        elif format_type.upper() == 'PNG':
            converter.convert_to_image(output_path, 'PNG')
        elif format_type.upper() == 'SVG':
            converter.convert_to_svg(output_path)
        elif format_type.upper() == 'PDF':
            converter.convert_to_pdf(output_path)
        elif format_type.upper() == 'DOCX':
            converter.convert_to_docx(output_path)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
        
        return jsonify({
            'success': True,
            'download_url': f'/download/{output_filename}',
            'filename': output_filename
        })
        
    except Exception as e:
        return jsonify({'error': f'Conversion failed: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return "File not found", 404
    except Exception as e:
        return f"Error downloading file: {str(e)}", 500

# Cleanup old files (run this periodically in production)
@app.route('/cleanup')
def cleanup_files():
    try:
        # Remove files older than 1 hour
        import time
        current_time = time.time()
        
        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                if os.path.isfile(file_path):
                    if current_time - os.path.getctime(file_path) > 3600:  # 1 hour
                        os.remove(file_path)
        
        return jsonify({'message': 'Cleanup completed'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8000)